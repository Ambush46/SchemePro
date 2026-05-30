"""
CONTROLLER: Scheme of Work
Generates scheme rows entirely from the curriculum database using
SchemeGenerator (app/services/scheme_engine.py).
No external API — fully deterministic, backend-only logic.

POST /scheme/generate  — build rows from DB + engine
POST /scheme/download  — deduct wallet, stream PDF / DOCX / ZIP
"""
import io
import json
import zipfile
from datetime import datetime, timezone
from flask import Blueprint, request, jsonify, send_file
from flask_login import login_required, current_user
from app import db
from app.models.wallet import Wallet, Payment
from app.models.pricing import DocumentPricing, GeneratedDocument
from app.models.subject import Subject
from app.models.scheme_draft import SchemeDraft
from app.services.scheme_engine import SchemeGenerator

scheme_bp = Blueprint('scheme', __name__)


# ── GENERATE ──────────────────────────────────────────────────────
@scheme_bp.route('/generate', methods=['POST'])
@login_required
def generate():
    """
    POST /scheme/generate
    Body:
      subject_id, grade, term, lessons_per_week, weeks, start_week,
      double_lesson, start_topic_id, start_subtopic_id,
      end_topic_id, end_subtopic_id,
      breaks: [{week, type, whole_week, start_lesson}]
    Returns:
      { success, rows, doc_id, references, curriculum_system }
    """
    data = request.get_json()
    current_user.record_activity('scheme_generator')

    subject_id = data.get('subject_id')
    if not subject_id:
        return jsonify({'success': False, 'error': 'subject_id is required.'}), 400

    subject = db.session.get(Subject, subject_id)
    if not subject:
        return jsonify({'success': False, 'error': 'Subject not found.'}), 404

    try:
        engine = SchemeGenerator(
            subject_id=int(subject_id),
            grade=data.get('grade', ''),
            term=int(data.get('term', 1)),
            lessons_per_week=int(data.get('lessons_per_week', 5)),
            weeks=int(data.get('weeks', 12)),
            start_week=int(data.get('start_week', 1)),
            double_lesson=data.get('double_lesson', ''),
            start_topic_id=data.get('start_topic_id'),
            start_subtopic_id=data.get('start_subtopic_id'),
            end_topic_id=data.get('end_topic_id'),
            end_subtopic_id=data.get('end_subtopic_id'),
            breaks=data.get('breaks', []),
        )
        rows = engine.generate()
        references = engine.references()
    except Exception as e:
        return jsonify({'success': False, 'error': f'Generation error: {str(e)}'}), 500

    if not rows:
        return jsonify({
            'success': False,
            'error': (
                'No curriculum content found for the selected topic range. '
                'Please ask your admin to add subtopics and content for this subject.'
            )
        }), 422

    # Log the generation
    doc = GeneratedDocument(
        user_id=current_user.id,
        subject_id=subject.id,
        subject_name=subject.name,
        grade=data.get('grade', ''),
        term=int(data.get('term', 1)),
        curriculum_system=subject.curriculum_system,
    )
    db.session.add(doc)
    db.session.commit()

    return jsonify({
        'success': True,
        'rows': rows,
        'doc_id': doc.id,
        'references': references,
        'curriculum_system': subject.curriculum_system,
    })


# ── DRAFTS ───────────────────────────────────────────────────────
@scheme_bp.route('/drafts', methods=['GET'])
@login_required
def get_draft():
    draft = SchemeDraft.query.filter_by(user_id=current_user.id).first()
    return jsonify({'success': True, 'draft': draft.to_dict() if draft else None})


@scheme_bp.route('/drafts', methods=['POST'])
@login_required
def save_draft():
    data = request.get_json(silent=True) or {}
    subject_id = data.get('subject_id')
    if not subject_id:
        return jsonify({'success': False, 'error': 'subject_id is required.'}), 400

    subject = db.session.get(Subject, subject_id)
    if not subject:
        return jsonify({'success': False, 'error': 'Subject not found.'}), 404

    payload = {
        'subject_id': int(subject_id),
        'step': int(data.get('step', 1)),
        'params': data.get('params') or {},
        'breaks': data.get('breaks') or [],
        'generated': data.get('generated') or [],
        'references': data.get('references') or [],
        'doc_id': data.get('doc_id'),
    }

    draft = SchemeDraft.query.filter_by(user_id=current_user.id).first()
    if draft:
        draft.subject_id = int(subject_id)
        draft.step = payload['step']
        draft.set_payload(payload)
    else:
        draft = SchemeDraft(
            user_id=current_user.id,
            subject_id=int(subject_id),
            step=payload['step'],
        )
        draft.set_payload(payload)
        db.session.add(draft)

    db.session.commit()
    return jsonify({'success': True, 'draft': draft.to_dict()})


# ── DOWNLOAD ──────────────────────────────────────────────────────
@scheme_bp.route('/download', methods=['POST'])
@login_required
def download():
    """
    POST /scheme/download
    Body: { doc_id, doc_type, rows, meta: {subject, grade, term,
            curriculum_system, references} }
    Deducts from wallet, returns binary file.
    """
    data = request.get_json()
    doc_type = data.get('doc_type', 'pdf')
    rows = data.get('rows', [])
    meta = data.get('meta', {})
    doc_id = data.get('doc_id')

    pricing = DocumentPricing.query.filter_by(doc_type=doc_type).first()
    if not pricing:
        return jsonify({'success': False, 'error': 'Invalid document type.'}), 400

    wallet = current_user.wallet
    if not wallet:
        wallet = Wallet(user_id=current_user.id, balance=0.0)
        db.session.add(wallet)
        db.session.commit()

    idempotency_key = data.get('idempotency_key')
    if idempotency_key:
        existing_payment = Payment.query.filter_by(user_id=current_user.id, idempotency_key=idempotency_key).first()
        if existing_payment:
            payment = existing_payment
        else:
            if wallet.balance < pricing.price:
                return jsonify({
                    'success': False,
                    'error': 'insufficient_balance',
                    'required': pricing.price,
                    'balance': wallet.balance,
                }), 402
            try:
                payment = wallet.debit(pricing.price, doc_type, idempotency_key=idempotency_key)
                db.session.commit()
            except ValueError as e:
                return jsonify({'success': False, 'error': str(e)}), 402
    else:
        if wallet.balance < pricing.price:
            return jsonify({
                'success': False,
                'error': 'insufficient_balance',
                'required': pricing.price,
                'balance': wallet.balance,
            }), 402
        try:
            payment = wallet.debit(pricing.price, doc_type)
            db.session.commit()
        except ValueError as e:
            return jsonify({'success': False, 'error': str(e)}), 402

    if doc_id:
        doc = db.session.get(GeneratedDocument, doc_id)
        if doc and doc.user_id == current_user.id:
            doc.doc_type = doc_type
            doc.downloaded_at = datetime.now(timezone.utc)
            doc.payment_ref = payment.transaction_id
            db.session.commit()

    is_cbc = meta.get('curriculum_system') == 'CBC'
    headers = (
        ['Wk', 'Lsn', 'Strand', 'Sub-Strand', 'Lesson Learning Outcomes',
         'Key Inquiry Questions', 'Learning Experiences', 'Learning Resources',
         'Assessment Methods', 'Refl']
        if is_cbc else
        ['Wk', 'Lsn', 'Topic', 'Sub-Topic', 'Objectives',
         'L/Activities', 'L/T Aids', 'Reference', 'Remarks']
    )
    keys = (
        ['wk', 'lsn', 'strand', 'substrand', 'outcomes', 'inquiry',
         'experiences', 'resources', 'assessment', 'refl']
        if is_cbc else
        ['wk', 'lsn', 'topic', 'subtopic', 'objectives',
         'activities', 'aids', 'reference', 'remarks']
    )

    title = f"{meta.get('subject','Subject')} — {meta.get('grade','')} Term {meta.get('term','')}"
    references = meta.get('references', [])

    if doc_type == 'pdf':
        buf = _build_pdf(title, meta, headers, keys, rows, is_cbc, references)
        return send_file(buf, mimetype='application/pdf',
                         as_attachment=True, download_name=f'{title}.pdf')

    elif doc_type == 'docx':
        buf = _build_docx(title, meta, headers, keys, rows, is_cbc, references)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True, download_name=f'{title}.docx'
        )

    elif doc_type == 'zip':
        pdf_buf = _build_pdf(title, meta, headers, keys, rows, is_cbc, references)
        docx_buf = _build_docx(title, meta, headers, keys, rows, is_cbc, references)
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f'{title}.pdf', pdf_buf.read())
            zf.writestr(f'{title}.docx', docx_buf.read())
            zf.writestr('README.txt', (
                f'SchemePro — {title}\n'
                f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n'
                f'Teacher: {current_user.name}\n\n'
                f'Files:\n'
                f'  {title}.pdf   — Print-ready PDF\n'
                f'  {title}.docx  — Editable Word document\n'
            ))
        zip_buf.seek(0)
        return send_file(zip_buf, mimetype='application/zip',
                         as_attachment=True, download_name=f'{title}.zip')

    return jsonify({'success': False, 'error': 'Unknown doc_type.'}), 400


# ── PDF BUILDER ────────────────────────────────────────────────────
def _build_pdf(title, meta, headers, keys, rows, is_cbc, references):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                    Paragraph, Spacer)
    from reportlab.lib.styles import ParagraphStyle

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=landscape(A4),
        rightMargin=1 * cm, leftMargin=1 * cm,
        topMargin=1.5 * cm, bottomMargin=1 * cm
    )
    blue = colors.HexColor('#1D4ED8')
    amber = colors.HexColor('#D97706')

    title_style = ParagraphStyle('T', fontSize=13, fontName='Helvetica-Bold',
                                 spaceAfter=3, textColor=blue)
    sub_style = ParagraphStyle('S', fontSize=9, fontName='Helvetica',
                               spaceAfter=2, textColor=colors.HexColor('#4B5563'))
    cell_style = ParagraphStyle('C', fontSize=7.5, fontName='Helvetica', leading=10)
    hdr_style = ParagraphStyle('H', fontSize=7, fontName='Helvetica-Bold',
                               textColor=colors.white)

    story = []
    subject = meta.get('subject', '')
    grade = meta.get('grade', '')
    term = meta.get('term', '')
    system = meta.get('curriculum_system', '844')

    if system == '844':
        story.append(Paragraph(f'{subject.upper()} SCHEME OF WORK', title_style))
        story.append(Paragraph(grade.upper(), sub_style))
        story.append(Paragraph(f'TERM {term}', sub_style))
        if references:
            story.append(Paragraph('REFERENCES:', sub_style))
            for r in references:
                story.append(Paragraph(f'• {r}', sub_style))
    else:
        story.append(Paragraph(
            f'{subject} {grade} — Term {term} Scheme of Work (CBC)', title_style))
        if references:
            story.append(Paragraph('REFERENCES: ' + ' | '.join(references), sub_style))
    story.append(Spacer(1, 8))

    n = len(headers)
    narrow = 1.1 * cm
    usable = 27.7 * cm - 2 * narrow
    col_widths = [narrow, narrow] + [usable / (n - 2)] * (n - 2)

    tdata = [[Paragraph(h, hdr_style) for h in headers]]
    for row in rows:
        tdata.append([Paragraph(str(row.get(k, '')), cell_style) for k in keys])

    style_cmds = [
        ('BACKGROUND', (0, 0), (-1, 0), blue),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTSIZE', (0, 0), (-1, 0), 7),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1),
         [colors.white, colors.HexColor('#EFF6FF')]),
        ('GRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#D1D5DB')),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
    ]
    for i, row in enumerate(rows, start=1):
        if 'BREAK' in str(row.get('topic', row.get('substrand', ''))).upper():
            style_cmds.append(
                ('BACKGROUND', (0, i), (-1, i), colors.HexColor('#FFF7ED')))
            style_cmds.append(
                ('TEXTCOLOR', (0, i), (-1, i), amber))

    t = Table(tdata, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle(style_cmds))
    story.append(t)
    doc.build(story)
    buf.seek(0)
    return buf


# ── DOCX BUILDER ───────────────────────────────────────────────────
def _build_docx(title, meta, headers, keys, rows, is_cbc, references):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Cm(1.5)
    section.top_margin = section.bottom_margin = Cm(1.5)

    subject = meta.get('subject', '')
    grade = meta.get('grade', '')
    term = meta.get('term', '')
    system = meta.get('curriculum_system', '844')

    def add_run(text, size=11, bold=False, color=(0x1D, 0x4E, 0xD8)):
        p = document.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(*color)
        p.paragraph_format.space_after = Pt(1)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    if system == '844':
        add_run(f'{subject.upper()} SCHEME OF WORK', 13, True)
        add_run(grade.upper(), 11, True)
        add_run(f'TERM {term}', 11, True)
        if references:
            add_run('REFERENCES:', 9, True, (0x4B, 0x55, 0x63))
            for r in references:
                add_run(f'  • {r}', 9, False, (0x4B, 0x55, 0x63))
    else:
        add_run(f'{subject} {grade} — Term {term} Scheme of Work (CBC)', 13, True)
        if references:
            add_run('REFERENCES: ' + ' | '.join(references), 9, False, (0x4B, 0x55, 0x63))

    document.add_paragraph()

    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, '1D4ED8')

    for row_data in rows:
        is_break = 'BREAK' in str(
            row_data.get('topic', row_data.get('substrand', ''))).upper()
        row = table.add_row()
        for i, key in enumerate(keys):
            cell = row.cells[i]
            cell.text = str(row_data.get(key, ''))
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7.5)
                    if is_break:
                        run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
            if is_break:
                set_cell_bg(cell, 'FFF7ED')

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf
